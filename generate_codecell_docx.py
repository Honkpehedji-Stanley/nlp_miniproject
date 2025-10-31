"""
Script de generation d'un document explicatif detaille pour detection_intention.ipynb
Explique cellule par cellule (seulement code), ligne par ligne
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    """Ajoute un saut de page."""
    doc.add_page_break()

def add_styled_heading(doc, text, level):
    """Ajoute un titre avec style."""
    heading = doc.add_heading(text, level=level)
    heading.runs[0].font.name = 'Arial'
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    return heading

def add_styled_paragraph(doc, text, bold=False, italic=False, color=None):
    """Ajoute un paragraphe avec style."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return para

def add_code_block(doc, code):
    """Ajoute un bloc de code avec style."""
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    # Fond gris pour le code
    para_format = para.paragraph_format
    para_format.left_indent = Inches(0.5)
    para_format.space_before = Pt(6)
    para_format.space_after = Pt(6)
    
    # Bordure
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '1')
        border.set(qn('w:color'), '808080')
        pBdr.append(border)
    pPr.append(pBdr)
    
    return para

def add_info_box(doc, title, text):
    """Ajoute une boite d'information."""
    para = doc.add_paragraph()
    
    # Titre de la boite
    title_run = para.add_run(f"💡 {title}\n")
    title_run.bold = True
    title_run.font.size = Pt(11)
    title_run.font.color.rgb = RGBColor(0, 102, 204)
    
    # Contenu
    content_run = para.add_run(text)
    content_run.font.size = Pt(10)
    content_run.font.color.rgb = RGBColor(60, 60, 60)
    
    # Style de la boite
    para_format = para.paragraph_format
    para_format.left_indent = Inches(0.3)
    para_format.space_before = Pt(6)
    para_format.space_after = Pt(6)
    
    return para

def add_bullet_point(doc, text):
    """Ajoute un point de liste."""
    para = doc.add_paragraph(text, style='List Bullet')
    para.paragraph_format.left_indent = Inches(0.25)
    return para

def create_documentation():
    """Cree la documentation complete cellule par cellule."""
    doc = Document()
    
    # Configuration des marges
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # PAGE DE TITRE
    title = doc.add_heading('EXPLICATION DETAILLEE', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(18)
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_heading('Notebook detection_intention.ipynb', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    
    project_info = doc.add_paragraph()
    project_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = project_info.add_run(
        'Explication Cellule par Cellule (Code)\n'
        'Ligne par Ligne avec Justifications\n\n'
    )
    run.font.size = Pt(12)
    run.italic = True
    
    doc.add_paragraph()
    
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(
        'Classification d\'Intentions avec CamemBERT\n'
        'Octobre 2025\n'
    )
    info_run.font.size = Pt(11)
    
    add_page_break(doc)
    
    # TABLE DES MATIERES
    add_styled_heading(doc, 'Table des Matieres', level=1)
    
    toc_items = [
        'Cellule 1: Verification GPU et Installation des Dependances',
        'Cellule 2: Montage Google Drive et Configuration',
        'Cellule 3: Chargement et Analyse des Datasets',
        'Cellule 4: Preprocessing et Parsing des Entities',
        'Cellule 5: Preparation des Datasets HuggingFace',
        'Cellule 6: Fine-Tuning Intent Classification',
        'Cellule 7: Evaluation et Post-Processing',
        'Cellule 8: Preparation Dataset NER',
        'Cellule 9: Fine-Tuning NER',
        'Cellule 10: Evaluation NER',
        'Cellule 11: Pipeline d\'Inference Complet',
    ]
    
    for i, item in enumerate(toc_items, 1):
        doc.add_paragraph(f'{i}. {item}', style='List Number')
    
    add_page_break(doc)
    
    # CELLULE 1: VERIFICATION GPU
    add_styled_heading(doc, 'Cellule 1: Verification GPU et Installation des Dependances', level=1)
    
    add_styled_paragraph(doc, 'Objectif de la cellule:', bold=True)
    add_styled_paragraph(doc, 
        'Cette cellule initialise l\'environnement d\'execution en verifiant la disponibilite '
        'du GPU et en installant les packages necessaires pour le deep learning avec Transformers.')
    
    add_styled_heading(doc, 'Code complet:', level=2)
    add_code_block(doc, 
        'import torch\n'
        '\n'
        'print(\'=\'*70)\n'
        'print(\'DETECTION INTENTION - Transformers + CamemBERT\')\n'
        'print(\'=\'*70)\n'
        '\n'
        'print(f\'\\nPyTorch version: {torch.__version__}\')\n'
        'print(f\'CUDA disponible: {torch.cuda.is_available()}\')\n'
        '\n'
        'if torch.cuda.is_available():\n'
        '    print(f\'GPU: {torch.cuda.get_device_name(0)}\')\n'
        '    print(f\'Memoire GPU: {torch.cuda.get_device_properties(0).total_memory / 1e9:.2f} GB\')\n'
        'else:\n'
        '    print(\'ATTENTION: Pas de GPU detecte!\')\n'
        '    print(\'Allez dans Runtime > Change runtime type > GPU\')\n'
        '\n'
        '!pip install -q transformers datasets evaluate seqeval accelerate scikit-learn langdetect\n'
        'print(\'\\nInstallation terminee!\')'
    )
    
    add_styled_heading(doc, 'Explication ligne par ligne:', level=2)
    
    add_styled_paragraph(doc, 'Ligne 1: import torch', bold=True)
    add_styled_paragraph(doc, 
        'Importe PyTorch, la bibliotheque de deep learning qui gere les calculs sur GPU/CPU. '
        'PyTorch est le framework sous-jacent utilise par Transformers de HuggingFace.')
    
    add_info_box(doc, 'Pourquoi PyTorch ?',
        'PyTorch offre une API flexible et pythonique, un support GPU excellent, et une '
        'integration native avec HuggingFace Transformers. Alternative: TensorFlow/Keras.')
    
    add_styled_paragraph(doc, 'Lignes 3-5: Affichage du titre', bold=True)
    add_styled_paragraph(doc, 
        'Affiche un titre encadre pour identifier clairement le debut de l\'execution. '
        'Les \'=\'*70 creent une ligne de 70 caracteres egale pour un affichage structure.')
    
    add_styled_paragraph(doc, 'Ligne 7: print(f\'PyTorch version: {torch.__version__}\')', bold=True)
    add_styled_paragraph(doc, 
        'Affiche la version de PyTorch installee. Important pour la reproductibilite : certaines '
        'fonctionnalites varient entre versions. Version recommandee : 2.0+')
    
    add_styled_paragraph(doc, 'Ligne 8: print(f\'CUDA disponible: {torch.cuda.is_available()}\')', bold=True)
    add_styled_paragraph(doc, 
        'Verifie si CUDA (Compute Unified Device Architecture de NVIDIA) est disponible. '
        'CUDA permet d\'utiliser le GPU pour accelerer les calculs. Retourne True si GPU detecte, False sinon.')
    
    add_info_box(doc, 'Importance du GPU',
        'L\'entrainement de Transformers est 10-50x plus rapide sur GPU que sur CPU. '
        'CamemBERT-base a 110M parametres : 6 epochs sur CPU = plusieurs jours, sur GPU = 30-60 minutes.')
    
    add_styled_paragraph(doc, 'Lignes 10-12: Informations GPU si disponible', bold=True)
    add_styled_paragraph(doc, 
        'Si CUDA est disponible, affiche le nom du GPU (ex: "Tesla T4", "V100") et sa memoire '
        'totale en GB. Ces informations permettent d\'ajuster la taille des batchs pour eviter '
        'les erreurs "Out of Memory".')
    
    add_code_block(doc, 
        'torch.cuda.get_device_name(0)  # Nom du premier GPU (index 0)\n'
        'torch.cuda.get_device_properties(0).total_memory  # Memoire en bytes\n'
        '/ 1e9  # Conversion bytes -> GB (divise par 1 milliard)'
    )
    
    add_styled_paragraph(doc, 'Lignes 13-15: Avertissement si pas de GPU', bold=True)
    add_styled_paragraph(doc, 
        'Si aucun GPU n\'est detecte, affiche un message d\'avertissement indiquant comment '
        'activer le GPU dans Google Colab. L\'execution reste possible sur CPU mais sera tres lente.')
    
    add_styled_paragraph(doc, 'Ligne 17: !pip install -q transformers datasets evaluate seqeval accelerate scikit-learn langdetect', bold=True)
    add_styled_paragraph(doc, 
        'Installe les packages Python necessaires. Le ! indique une commande shell dans Jupyter/Colab. '
        'L\'option -q (quiet) reduit les messages d\'installation.')
    
    add_styled_paragraph(doc, 'Packages installes:', bold=True)
    add_bullet_point(doc, 'transformers : Bibliotheque HuggingFace pour les modeles pre-entraines (CamemBERT, BERT, GPT)')
    add_bullet_point(doc, 'datasets : Gestion efficace des datasets (loading, preprocessing, caching)')
    add_bullet_point(doc, 'evaluate : Metriques d\'evaluation (accuracy, F1, precision, recall)')
    add_bullet_point(doc, 'seqeval : Metriques specialisees pour le NER (respecte le format BIO)')
    add_bullet_point(doc, 'accelerate : Gestion automatique GPU/CPU et mixed precision training')
    add_bullet_point(doc, 'scikit-learn : Calcul de class weights, metriques, preprocessing')
    add_bullet_point(doc, 'langdetect : Detection automatique de la langue d\'un texte')
    
    add_info_box(doc, 'Choix des versions',
        'Les versions sont geres automatiquement par pip. Pour la reproductibilite en production, '
        'utiliser requirements.txt avec versions figees (ex: transformers==4.35.0).')
    
    add_page_break(doc)
    
    # CELLULE 2: GOOGLE DRIVE
    add_styled_heading(doc, 'Cellule 2: Montage Google Drive et Configuration', level=1)
    
    add_styled_paragraph(doc, 'Objectif de la cellule:', bold=True)
    add_styled_paragraph(doc, 
        'Monte Google Drive pour acceder aux datasets et configure l\'environnement pour '
        'eviter les outils de tracking non necessaires.')
    
    add_styled_heading(doc, 'Code complet:', level=2)
    add_code_block(doc, 
        'from google.colab import drive\n'
        'import os\n'
        '\n'
        '# Monter Google Drive\n'
        'drive.mount(\'/content/drive\')\n'
        '\n'
        '# Desactiver WandB\n'
        'os.environ[\'WANDB_DISABLED\'] = \'true\'\n'
        '\n'
        '# Chemins\n'
        'workdir = \'/content/drive/MyDrive/dataset\'\n'
        'os.makedirs(workdir, exist_ok=True)\n'
        '\n'
        'print(\'Working directory:\', workdir)\n'
        'print(\'WandB: DESACTIVE\')'
    )
    
    add_styled_heading(doc, 'Explication ligne par ligne:', level=2)
    
    add_styled_paragraph(doc, 'Ligne 1: from google.colab import drive', bold=True)
    add_styled_paragraph(doc, 
        'Importe le module drive de google.colab qui permet de monter Google Drive comme un '
        'systeme de fichiers dans l\'environnement Colab. Ceci est specifique a Google Colab.')
    
    add_styled_paragraph(doc, 'Ligne 2: import os', bold=True)
    add_styled_paragraph(doc, 
        'Importe le module os (operating system) qui fournit des fonctions pour interagir avec '
        'le systeme d\'exploitation : manipulation de chemins, variables d\'environnement, creation de dossiers.')
    
    add_styled_paragraph(doc, 'Ligne 5: drive.mount(\'/content/drive\')', bold=True)
    add_styled_paragraph(doc, 
        'Monte Google Drive sur le chemin /content/drive. Cette commande affiche un lien '
        'd\'authentification : l\'utilisateur doit se connecter et autoriser l\'acces. '
        'Apres montage, les fichiers Drive sont accessibles comme des fichiers locaux.')
    
    add_info_box(doc, 'Pourquoi Google Drive ?',
        'Les fichiers dans l\'environnement Colab sont temporaires et supprimes apres deconnexion. '
        'Google Drive permet de :\n'
        '- Stocker les datasets de maniere persistante\n'
        '- Sauvegarder les modeles entraines\n'
        '- Partager les donnees entre sessions')
    
    add_styled_paragraph(doc, 'Ligne 8: os.environ[\'WANDB_DISABLED\'] = \'true\'', bold=True)
    add_styled_paragraph(doc, 
        'Desactive Weights & Biases (WandB), un outil de tracking d\'experimentations. '
        'os.environ modifie les variables d\'environnement. WandB est utile pour le suivi '
        'professionnel mais non necessaire pour ce projet.')
    
    add_info_box(doc, 'Qu\'est-ce que WandB ?',
        'Weights & Biases est une plateforme qui enregistre automatiquement les metriques, '
        'hyperparametres et graphiques pendant l\'entrainement. Desactive ici car :\n'
        '- Necessite une connexion internet stable\n'
        '- Ajoute de la latence\n'
        '- Pas necessaire pour un projet educatif')
    
    add_styled_paragraph(doc, 'Ligne 11: workdir = \'/content/drive/MyDrive/dataset\'', bold=True)
    add_styled_paragraph(doc, 
        'Define le repertoire de travail ou sont stockes les datasets et ou seront sauvegardes '
        'les modeles. /content/drive/MyDrive/ correspond a "Mon Drive" dans l\'interface Google Drive.')
    
    add_styled_paragraph(doc, 'Ligne 12: os.makedirs(workdir, exist_ok=True)', bold=True)
    add_styled_paragraph(doc, 
        'Cree le dossier workdir s\'il n\'existe pas. exist_ok=True evite une erreur si le '
        'dossier existe deja. Cette ligne assure que le repertoire est disponible pour la suite.')
    
    add_code_block(doc, 
        '# Sans exist_ok=True :\n'
        'os.makedirs(\'/mon/dossier\')  # OK la premiere fois\n'
        'os.makedirs(\'/mon/dossier\')  # ERREUR: le dossier existe deja!\n'
        '\n'
        '# Avec exist_ok=True :\n'
        'os.makedirs(\'/mon/dossier\', exist_ok=True)  # OK meme si existe'
    )
    
    add_styled_paragraph(doc, 'Lignes 14-15: Affichage des informations', bold=True)
    add_styled_paragraph(doc, 
        'Confirme le chemin de travail et l\'etat de WandB. Important pour verifier que la '
        'configuration est correcte avant de continuer.')
    
    add_page_break(doc)
    
    # CELLULE 3: CHARGEMENT DATASETS
    add_styled_heading(doc, 'Cellule 3: Chargement et Analyse des Datasets', level=1)
    
    add_styled_paragraph(doc, 'Objectif de la cellule:', bold=True)
    add_styled_paragraph(doc, 
        'Charge les fichiers CSV contenant les donnees d\'entrainement et de test, puis affiche '
        'des statistiques pour comprendre la distribution des classes.')
    
    add_styled_heading(doc, 'Code complet:', level=2)
    add_code_block(doc, 
        'import pandas as pd\n'
        '\n'
        'train_path = os.path.join(workdir, \'train_set.csv\')\n'
        'test_path = os.path.join(workdir, \'test_set.csv\')\n'
        '\n'
        '# Verification existence\n'
        'for p in [train_path, test_path]:\n'
        '    if not os.path.exists(p):\n'
        '        print(f\'ERREUR: Fichier non trouve: {p}\')\n'
        '        raise FileNotFoundError(p)\n'
        '\n'
        '# Chargement\n'
        'train_df = pd.read_csv(train_path, encoding=\'utf-8\')\n'
        'test_df = pd.read_csv(test_path, encoding=\'utf-8\')\n'
        '\n'
        'print(f\'Train shape: {train_df.shape}\')\n'
        'print(f\'Test shape: {test_df.shape}\')\n'
        '\n'
        'print(\'\\nDistribution des classes (Train):\')\n'
        'print(train_df[\'intent\'].value_counts())\n'
        'print(f\'\\nPourcentages:\')\n'
        'print(train_df[\'intent\'].value_counts(normalize=True).apply(lambda x: f\'{x:.2%}\'))\n'
        '\n'
        'display(train_df.head(3))'
    )
    
    add_styled_heading(doc, 'Explication ligne par ligne:', level=2)
    
    add_styled_paragraph(doc, 'Ligne 1: import pandas as pd', bold=True)
    add_styled_paragraph(doc, 
        'Importe pandas, LA bibliotheque Python pour la manipulation de donnees tabulaires. '
        'Pandas fournit des structures DataFrame (tableau 2D) et Series (colonne) avec des '
        'operations optimisees.')
    
    add_info_box(doc, 'Pourquoi pandas ?',
        'Pandas est le standard pour le data science en Python :\n'
        '- Lecture facile de CSV, Excel, JSON, SQL\n'
        '- Operations vectorisees (rapides)\n'
        '- Gestion automatique des types de donnees\n'
        '- Integration avec numpy, sklearn, matplotlib')
    
    add_styled_paragraph(doc, 'Lignes 3-4: Construction des chemins', bold=True)
    add_styled_paragraph(doc, 
        'Construit les chemins complets vers les fichiers CSV en utilisant os.path.join(). '
        'Cette fonction gere automatiquement les separateurs de chemin selon l\'OS (/ sur Linux/Mac, \\ sur Windows).')
    
    add_code_block(doc, 
        '# Mauvaise pratique (fragile) :\n'
        'train_path = workdir + \'/train_set.csv\'  # Probleme si workdir se termine deja par /\n'
        '\n'
        '# Bonne pratique (robuste) :\n'
        'train_path = os.path.join(workdir, \'train_set.csv\')  # Toujours correct'
    )
    
    add_styled_paragraph(doc, 'Lignes 7-10: Verification d\'existence', bold=True)
    add_styled_paragraph(doc, 
        'Verifie que les deux fichiers existent avant de tenter de les charger. Cette verification '
        'preventive evite des erreurs cryptiques et affiche un message clair si un fichier manque.')
    
    add_styled_paragraph(doc, 'Decomposition de la boucle:', italic=True)
    add_bullet_point(doc, 'for p in [train_path, test_path]: Itere sur les deux chemins')
    add_bullet_point(doc, 'os.path.exists(p): Retourne True si le fichier existe, False sinon')
    add_bullet_point(doc, 'raise FileNotFoundError(p): Leve une exception qui arrete l\'execution avec un message d\'erreur')
    
    add_info_box(doc, 'Importance de la validation',
        'Verifier les fichiers des le debut evite de perdre du temps d\'execution. Sans cette '
        'verification, l\'erreur apparaitrait seulement a la ligne 13, apres avoir deja execute '
        'du code inutilement.')
    
    add_styled_paragraph(doc, 'Lignes 13-14: Chargement des CSV', bold=True)
    add_styled_paragraph(doc, 
        'Charge les fichiers CSV dans des DataFrames pandas. encoding=\'utf-8\' assure que les '
        'caracteres francais (e, a, c, etc.) sont correctement interpretes.')
    
    add_code_block(doc, 
        '# pd.read_csv() fait automatiquement :\n'
        '# 1. Detection des separateurs (, ; tab)\n'
        '# 2. Inference des types de colonnes (int, float, string)\n'
        '# 3. Creation d\'un index numerique (0, 1, 2, ...)\n'
        '# 4. Optimisation memoire'
    )
    
    add_styled_paragraph(doc, 'Lignes 16-17: Affichage des dimensions', bold=True)
    add_styled_paragraph(doc, 
        'Affiche la forme des DataFrames avec .shape qui retourne un tuple (nombre_lignes, nombre_colonnes). '
        'Permet de verifier rapidement si le chargement a reussi et si les dimensions sont attendues.')
    
    add_code_block(doc, 
        '# Exemple de sortie :\n'
        'Train shape: (8000, 3)  # 8000 lignes, 3 colonnes\n'
        'Test shape: (2000, 3)   # 2000 lignes, 3 colonnes'
    )
    
    add_styled_paragraph(doc, 'Lignes 19-20: Distribution des classes', bold=True)
    add_styled_paragraph(doc, 
        'Affiche le nombre d\'exemples par classe d\'intention avec value_counts(). Cette methode '
        'compte les occurrences uniques dans la colonne \'intent\' et les trie par ordre decroissant.')
    
    add_code_block(doc, 
        '# Exemple de sortie :\n'
        'TRIP          3200\n'
        'NOT_TRIP      2400\n'
        'NOT_FRENCH    1600\n'
        'UNKNOWN        800\n'
        'Name: intent, dtype: int64'
    )
    
    add_info_box(doc, 'Importance de la distribution',
        'Comprendre la distribution des classes est crucial pour :\n'
        '- Detecter un desequilibre (qui necessite class weighting)\n'
        '- Interpreter les metriques (accuracy est trompeuse si desequilibre)\n'
        '- Ajuster la strategie d\'echantillonnage')
    
    add_styled_paragraph(doc, 'Lignes 21-22: Pourcentages', bold=True)
    add_styled_paragraph(doc, 
        'Affiche la meme distribution en pourcentages avec normalize=True. La fonction lambda '
        'formate chaque valeur en pourcentage avec 2 decimales.')
    
    add_code_block(doc, 
        'value_counts(normalize=True)  # Divise par le total -> valeurs entre 0 et 1\n'
        '.apply(lambda x: f\'{x:.2%}\')  # Formate 0.4 -> \'40.00%\'\n'
        '\n'
        '# Exemple de sortie :\n'
        'TRIP          40.00%\n'
        'NOT_TRIP      30.00%\n'
        'NOT_FRENCH    20.00%\n'
        'UNKNOWN       10.00%'
    )
    
    add_styled_paragraph(doc, 'Ligne 24: display(train_df.head(3))', bold=True)
    add_styled_paragraph(doc, 
        'Affiche les 3 premieres lignes du DataFrame d\'entrainement dans un format interactif '
        '(specifique a Jupyter/Colab). Permet de voir la structure des donnees : noms de colonnes, '
        'types, exemples de valeurs.')
    
    add_info_box(doc, 'display() vs print()',
        'display() cree un tableau HTML formate et interactif dans Jupyter/Colab. '
        'print() affiche du texte brut moins lisible. Toujours utiliser display() pour les DataFrames '
        'dans les notebooks.')
    
    add_page_break(doc)
    
    # CELLULE 4: PREPROCESSING
    add_styled_heading(doc, 'Cellule 4: Preprocessing et Parsing des Entities', level=1)
    
    add_styled_paragraph(doc, 'Objectif de la cellule:', bold=True)
    add_styled_paragraph(doc, 
        'Nettoie les donnees en retirant les valeurs manquantes et parse les annotations JSON '
        'des entites (villes de depart/destination) en validant leur coherence.')
    
    add_styled_heading(doc, 'Code complet:', level=2)
    add_code_block(doc, 
        'import json\n'
        '\n'
        '# Nettoyer les colonnes necessaires\n'
        'train_df = train_df[[\'text\', \'intent\', \'entities\']].dropna(subset=[\'text\', \'intent\']).reset_index(drop=True)\n'
        'test_df = test_df[[\'text\', \'intent\', \'entities\']].dropna(subset=[\'text\', \'intent\']).reset_index(drop=True)\n'
        '\n'
        'def parse_entities_field(row):\n'
        '    """Parse la colonne entities (JSON) et valide."""\n'
        '    try:\n'
        '        ents = json.loads(row[\'entities\']) if pd.notna(row[\'entities\']) else []\n'
        '    except:\n'
        '        ents = []\n'
        '\n'
        '    valid = []\n'
        '    txt = row.get(\'text\', \'\')\n'
        '    for ent in ents:\n'
        '        if isinstance(ent, dict) and \'start\' in ent and \'end\' in ent and \'label\' in ent:\n'
        '            if 0 <= ent[\'start\'] < ent[\'end\'] <= len(txt):\n'
        '                valid.append(ent)\n'
        '    return valid\n'
        '\n'
        '# Parser les entities\n'
        'train_df[\'parsed_entities\'] = train_df.apply(parse_entities_field, axis=1)\n'
        'test_df[\'parsed_entities\'] = test_df.apply(parse_entities_field, axis=1)\n'
        '\n'
        'print(\'Entities parsees avec succes!\')\n'
        'print(f\'\\nExemple TRIP avec entities:\')\n'
        'trip_ex = train_df[train_df[\'intent\'] == \'TRIP\'].iloc[0]\n'
        'print(f\'  Texte: {trip_ex["text"]}\')\n'
        'print(f\'  Entities: {trip_ex["parsed_entities"]}\')'
    )
    
    add_styled_heading(doc, 'Explication ligne par ligne:', level=2)
    
    add_styled_paragraph(doc, 'Ligne 1: import json', bold=True)
    add_styled_paragraph(doc, 
        'Importe le module json pour parser les chaines JSON. JSON (JavaScript Object Notation) '
        'est un format texte pour representer des structures de donnees (listes, dictionnaires).')
    
    add_styled_paragraph(doc, 'Lignes 4-5: Nettoyage des colonnes', bold=True)
    add_styled_paragraph(doc, 
        'Cette ligne complexe effectue 3 operations en chaine :')
    
    add_styled_paragraph(doc, '1. Selection de colonnes: [[\'text\', \'intent\', \'entities\']]', bold=True)
    add_styled_paragraph(doc, 
        'Selectionne uniquement les 3 colonnes necessaires. Les doubles crochets [[]] indiquent '
        'qu\'on veut un DataFrame (pas une Serie). Cela elimine les colonnes inutiles.')
    
    add_styled_paragraph(doc, '2. Suppression valeurs manquantes: .dropna(subset=[\'text\', \'intent\'])', bold=True)
    add_styled_paragraph(doc, 
        'Supprime les lignes ou \'text\' ou \'intent\' est NaN (Not a Number = valeur manquante). '
        'subset precise quelles colonnes verifier. La colonne \'entities\' peut etre NaN (NOT_TRIP n\'a pas d\'entites).')
    
    add_info_box(doc, 'Pourquoi dropna ?',
        'Un texte manquant ou une intention manquante rend l\'exemple inutilisable pour l\'entrainement. '
        'Mieux vaut supprimer ces lignes que de risquer des erreurs plus tard.')
    
    add_styled_paragraph(doc, '3. Reset de l\'index: .reset_index(drop=True)', bold=True)
    add_styled_paragraph(doc, 
        'Apres dropna(), les indices sont discontinus (ex: 0, 1, 5, 8...). reset_index() recree '
        'un index sequentiel (0, 1, 2, 3...). drop=True supprime l\'ancien index au lieu de le garder '
        'comme colonne.')
    
    add_code_block(doc, 
        '# Avant reset_index :\n'
        'Index: 0, 1, 3, 5, 7  # Indices discontinus apres dropna()\n'
        '\n'
        '# Apres reset_index(drop=True) :\n'
        'Index: 0, 1, 2, 3, 4  # Indices continus'
    )
    
    add_styled_paragraph(doc, 'Lignes 7-20: Fonction parse_entities_field', bold=True)
    add_styled_paragraph(doc, 
        'Cette fonction complexe parse et valide les annotations d\'entites. Elle prend une ligne '
        'du DataFrame (row) et retourne une liste d\'entites valides.')
    
    add_styled_paragraph(doc, 'Lignes 9-12: Parsing JSON avec gestion d\'erreurs', bold=True)
    add_code_block(doc, 
        'try:\n'
        '    ents = json.loads(row[\'entities\']) if pd.notna(row[\'entities\']) else []\n'
        'except:\n'
        '    ents = []'
    )
    
    add_styled_paragraph(doc, 'Decomposition:', italic=True)
    add_bullet_point(doc, 'pd.notna(row[\'entities\']): Verifie que la valeur n\'est pas NaN')
    add_bullet_point(doc, 'json.loads(): Convertit une chaine JSON en objet Python (liste de dicts)')
    add_bullet_point(doc, 'try/except: Si le JSON est malformé, retourne liste vide au lieu de crasher')
    
    add_code_block(doc, 
        '# Exemple de valeur entities :\n'
        '\'[{"start": 3, "end": 8, "label": "Departure"}, {"start": 11, "end": 16, "label": "Destination"}]\'\n'
        '\n'
        '# Apres json.loads() :\n'
        '[\n'
        '    {"start": 3, "end": 8, "label": "Departure"},\n'
        '    {"start": 11, "end": 16, "label": "Destination"}\n'
        ']'
    )
    
    add_styled_paragraph(doc, 'Lignes 14-19: Validation des entites', bold=True)
    add_styled_paragraph(doc, 
        'Filtre les entites pour ne garder que celles qui sont valides. Une entite valide doit :')
    
    add_bullet_point(doc, 'Etre un dictionnaire (isinstance(ent, dict))')
    add_bullet_point(doc, 'Avoir les cles \'start\', \'end\', \'label\'')
    add_bullet_point(doc, 'Avoir des indices coherents : 0 <= start < end <= longueur_texte')
    
    add_info_box(doc, 'Pourquoi valider ?',
        'Les annotations peuvent contenir des erreurs :\n'
        '- Indices negatifs\n'
        '- Indices hors du texte (end > len(text))\n'
        '- start >= end (span vide)\n'
        'Ces erreurs causeraient des bugs dans le NER. Mieux vaut les filtrer maintenant.')
    
    add_code_block(doc, 
        '# Exemple d\'entite invalide :\n'
        '{"start": 50, "end": 55, "label": "Departure"}  # Si len(text) = 40, hors limites!\n'
        '\n'
        '# Exemple d\'entite valide :\n'
        'text = "De Paris a Lyon"  # len = 15\n'
        '{"start": 3, "end": 8, "label": "Departure"}  # 3 < 8 <= 15, OK!'
    )
    
    add_styled_paragraph(doc, 'Lignes 23-24: Application de la fonction', bold=True)
    add_styled_paragraph(doc, 
        'Applique parse_entities_field a chaque ligne avec .apply(). axis=1 indique qu\'on passe '
        'chaque ligne (row) a la fonction. Le resultat est stocke dans une nouvelle colonne \'parsed_entities\'.')
    
    add_code_block(doc, 
        '# .apply() equivalent a :\n'
        'parsed_entities = []\n'
        'for index, row in train_df.iterrows():\n'
        '    result = parse_entities_field(row)\n'
        '    parsed_entities.append(result)\n'
        'train_df[\'parsed_entities\'] = parsed_entities\n'
        '\n'
        '# Mais .apply() est optimise et plus concis'
    )
    
    add_styled_paragraph(doc, 'Lignes 26-29: Affichage d\'exemple', bold=True)
    add_styled_paragraph(doc, 
        'Affiche un exemple de requete TRIP avec ses entites parsees pour verifier visuellement '
        'que le parsing fonctionne correctement.')
    
    add_styled_paragraph(doc, 'Decomposition:', italic=True)
    add_bullet_point(doc, 'train_df[train_df[\'intent\'] == \'TRIP\']: Filtre pour ne garder que les TRIP')
    add_bullet_point(doc, '.iloc[0]: Prend la premiere ligne du resultat')
    add_bullet_point(doc, 'trip_ex["text"] et trip_ex["parsed_entities"]: Accede aux colonnes')
    
    add_page_break(doc)
    
    # Sauvegarder le document
    output_path = '/home/stanley-honkpehedji/Téléchargements/nlp_miniproject/Detection_Intention_Code_Explanation.docx'
    doc.save(output_path)
    print(f'Document genere : {output_path}')
    print(f'Note: Document incomplet pour demonstration. Continuer avec les autres cellules...')

if __name__ == '__main__':
    create_documentation()
